import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfkit
import tempfile
import zipfile
from django.conf import settings

def clean_html(html_text):
    """Clean HTML content while preserving formatting"""
    if not html_text:
        return ""
    
    soup = BeautifulSoup(html_text, 'html.parser')
    
    # Pre-process: Remove empty tags and br tags
    for tag in soup.find_all():
        if len(tag.get_text(strip=True)) == 0 or tag.name == 'br':
            tag.decompose()
    
    # Initialize result text
    result = []
    processed_text = set() # Keep track of processed text to avoid duplicates
    
    # Process each tag and convert to text with proper spacing
    for element in soup.descendants:
        if isinstance(element, str) and element.strip():
            # Handle plain text
            text = element.strip()
            if text not in processed_text:
                result.append(text)
                processed_text.add(text)
        elif element.name in ['h1', 'h2', 'h3']:
            # Handle headers
            text = element.get_text().strip()
            if text not in processed_text:
                result.extend(['', '', text, ''])
                processed_text.add(text)
        elif element.name == 'p':
            # Handle paragraphs
            text = element.get_text().strip()
            if text and text not in processed_text:
                result.extend(['', text, ''])
                processed_text.add(text)
        elif element.name in ['ul', 'ol']:
            # Handle lists
            result.append('')
            for i, li in enumerate(element.find_all('li', recursive=False)):
                text = li.get_text().strip()
                if text and text not in processed_text:
                    prefix = f"{i+1}. " if element.name == 'ol' else "• "
                    result.append(f"{prefix}{text}")
                    processed_text.add(text)
            result.append('')
        elif element.name == 'blockquote':
            # Handle blockquotes
            text = element.get_text().strip()
            if text and text not in processed_text:
                lines = text.split('\n')
                result.extend([''] + [f"> {line.strip()}" for line in lines if line.strip()] + [''])
                processed_text.add(text)
        elif element.name == 'em':
            # Handle emphasis
            text = element.get_text().strip()
            if text and text not in processed_text:
                result.append(f"_{text}_")
                processed_text.add(text)
        elif element.name == 'strong':
            # Handle strong
            text = element.get_text().strip()
            if text and text not in processed_text:
                result.append(f"**{text}**")
                processed_text.add(text)
    
    # Join all parts with proper spacing
    text = '\n'.join(result)
    
    # Clean up multiple blank lines while preserving intentional spacing
    lines = text.split('\n')
    cleaned_lines = []
    prev_empty = False
    
    for line in lines:
        line = line.rstrip()
        is_empty = not line.strip()
        
        if is_empty and prev_empty:
            continue
        
        cleaned_lines.append(line)
        prev_empty = is_empty
    
    # Final cleanup
    text = '\n'.join(cleaned_lines).strip()
    text = text.replace('\r\n', '\n')
    
    return text

def create_pdf_document(pages, is_translation=False):
    """Create a PDF document from pages with continuous text"""
    direction = 'rtl' if 'rtl' in (pages[0].detect_translated_direction() if is_translation else pages[0].detect_direction()) else 'ltr'
    
    # Combine all text with section headers
    content = []
    for page in pages:
        text = page.translated_text if is_translation else page.original_text
        content.append(f"<h2>Page {page.number}</h2>\n{text}")
    
    combined_text = "\n".join(content)
    
    html_content = f"""
        <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    @page {{
                        margin: 2.5cm;
                        @top-center {{
                            content: counter(page);
                            font-family: Arial, sans-serif;
                            font-size: 10pt;
                            color: #6B7280;
                        }}
                    }}
                    body {{
                        font-family: Arial, sans-serif;
                        font-size: 12pt;
                        line-height: 1.8;
                        color: #1F2937;
                        text-align: justify;
                    }}
                    h2 {{
                        color: #374151;
                        margin-top: 2rem;
                        margin-bottom: 1rem;
                        page-break-after: avoid;
                    }}
                    p {{
                        margin-bottom: 1rem;
                        text-align: justify;
                        page-break-inside: avoid;
                    }}
                    blockquote {{
                        margin: 1rem 0;
                        padding-left: 1rem;
                        border-left: 4px solid #E5E7EB;
                        color: #4B5563;
                    }}
                    ul, ol {{
                        margin: 1rem 0;
                        padding-left: 2rem;
                    }}
                    li {{
                        margin-bottom: 0.5rem;
                    }}
                </style>
            </head>
            <body dir="{direction}">
                <div class="content">
                    {combined_text}
                </div>
            </body>
        </html>
    """
    
    options = {
        'page-size': 'A4',
        'margin-top': '2.5cm',
        'margin-right': '2.5cm',
        'margin-bottom': '2.5cm',
        'margin-left': '2.5cm',
        'encoding': 'UTF-8',
        'no-outline': None,
        'enable-local-file-access': None
    }
    
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
        pdfkit.from_string(html_content, tmp.name, options=options)
        return tmp.name

def create_word_document(pages, is_translation=False):
    """Create a Word document with continuous text"""
    doc = Document()
    
    # Set default font and style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    
    # Get direction from first page
    direction = 'rtl' if 'rtl' in (pages[0].detect_translated_direction() if is_translation else pages[0].detect_direction()) else 'ltr'
    
    for page in pages:
        # Add page number as heading
        heading = doc.add_heading(f'Page {page.number}', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT if direction == 'rtl' else WD_ALIGN_PARAGRAPH.LEFT
        heading.style.font.color.rgb = RGBColor(55, 65, 81)
        
        # Add content
        text = page.translated_text if is_translation else page.original_text
        text = clean_html(text)
        
        # Split content into paragraphs
        paragraphs = text.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if direction == 'rtl' else WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run(para_text.strip())
                run.font.name = 'Arial'
                
                if para_text.startswith('#'):
                    run.bold = True
                    run.font.size = Pt(14)
                elif para_text.startswith('>'):
                    paragraph.style = 'Quote'
    
    # Save document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        doc.save(tmp.name)
        return tmp.name

def create_text_document(pages, is_translation=False):
    """Create a text document with continuous text"""
    content = []
    
    for page in pages:
        text = page.translated_text if is_translation else page.original_text
        text = clean_html(text)
        
        content.append(f"\n\n{'='*20}\nPage {page.number}\n{'='*20}\n\n{text}")
    
    with tempfile.NamedTemporaryFile(suffix='.txt', mode='w', encoding='utf-8', delete=False) as tmp:
        tmp.write('\n'.join(content))
        return tmp.name
    
def create_zip_archive(book):
    """Create ZIP archive containing Word documents of both original and translated text"""
    pages = book.pages.all().order_by('number')
    
    with tempfile.NamedTemporaryFile(suffix='.zip', delete=False) as tmp_zip:
        with zipfile.ZipFile(tmp_zip.name, 'w') as zip_file:
            # Original Word document
            docx_path = create_word_document(pages, is_translation=False)
            zip_file.write(docx_path, f'original/book.docx')
            
            # Translation Word document
            docx_path_tr = create_word_document(pages, is_translation=True)
            zip_file.write(docx_path_tr, f'translation/book.docx')
            
            # Cleanup temporary files
            for path in [docx_path, docx_path_tr]:
                os.unlink(path)
                
        return tmp_zip.name